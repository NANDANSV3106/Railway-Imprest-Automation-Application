"""
Anchor strings the docx_engine generators rely on to locate paragraphs/rows,
and a pre-generation sanity check that confirms a template still contains them.
Centralized here so a template edit only needs updating in one place.
"""
import docx


ANNEX1_ANCHORS = [
    "No. Station Imprest", "Sub : Recoupment", "Vr. No.",
    "Accepted for Rs.", "For the month of", "Imprest Account Of",
    "Opening Balance", "Closing Balance", "Sanctioned Amount",
]
LABOUR_ANCHORS = ["LABOUR NAME:", "Total amount paid", "DATE:"]


def validate_template(template_path, required_anchors, min_tables=1):
    """
    Quick sanity check run before generation: confirms the template file
    opens and still contains the anchor text the fill logic searches for.
    Raises a specific, actionable error instead of letting generation
    fail deep inside table-fill logic with a blank field or IndexError.
    """
    try:
        doc = docx.Document(template_path)
    except Exception as e:
        raise RuntimeError(f"Could not open template '{template_path}': {e}") from e

    if len(doc.tables) < min_tables:
        raise RuntimeError(
            f"Template '{template_path}' has {len(doc.tables)} table(s), "
            f"expected at least {min_tables}. The template structure may have changed."
        )

    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text

    missing = [a for a in required_anchors if a not in all_text]
    if missing:
        raise RuntimeError(
            f"Template '{template_path}' is missing expected text: {', '.join(missing)}. "
            "The template may have been edited — the fill logic needs matching updates."
        )
    return doc

