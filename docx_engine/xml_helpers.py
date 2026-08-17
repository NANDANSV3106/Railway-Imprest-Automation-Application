"""
Low-level OOXML helpers for writing precisely-formatted Word runs/cells.
python-docx's high-level API cannot control every formatting detail
the reference documents need (exact spacing, run-level bold/size on
existing cell content, etc.) so these build/edit XML elements directly.
"""
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _make_run(text, bold=False, size=None, underline=False,
              font_ascii="Calibri", color=None):
    """Build a <w:r> element with optional formatting."""
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    if bold:
        rpr.append(OxmlElement('w:b'))
        rpr.append(OxmlElement('w:bCs'))
    if font_ascii:
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:ascii'), font_ascii)
        rf.set(qn('w:hAnsi'), font_ascii)
        rf.set(qn('w:cs'), font_ascii)
        rpr.append(rf)
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rpr.append(c)
    if size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size))
        rpr.append(sz)
        szcs = OxmlElement('w:szCs')
        szcs.set(qn('w:val'), str(size))
        rpr.append(szcs)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rpr.append(u)
    r.append(rpr)
    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    return r


def _clear_para_runs(para):
    """Remove all <w:r> children from a paragraph."""
    for r_elem in para._p.findall(qn('w:r')):
        para._p.remove(r_elem)


def _set_para_spacing(para, after="0", line="240", line_rule="auto"):
    """Add/replace <w:spacing> inside <w:pPr>."""
    ppr = para._p.find(qn('w:pPr'))
    if ppr is None:
        ppr = OxmlElement('w:pPr')
        para._p.insert(0, ppr)
    sp = ppr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        ppr.append(sp)
    sp.set(qn('w:after'), after)
    sp.set(qn('w:line'), line)
    sp.set(qn('w:lineRule'), line_rule)


def _set_cell_run(cell, text, bold=False, size=None, spacing_after="0",
                  spacing_line="240", font_ascii="Calibri", color="000000"):
    """
    Replace cell content with a single run.
    Applies spacing and font size to both pPr and the run itself
    (matching the exact XML pattern of the correct template).
    """
    para = cell.paragraphs[0]
    _clear_para_runs(para)
    _set_para_spacing(para, after=spacing_after, line=spacing_line)

    # Also write rPr into pPr (paragraph-mark formatting) so spacing is correct
    ppr = para._p.find(qn('w:pPr'))
    ppr_rpr = ppr.find(qn('w:rPr'))
    if ppr_rpr is None:
        ppr_rpr = OxmlElement('w:rPr')
        ppr.append(ppr_rpr)
    # Clear and rebuild pPr rPr
    for child in list(ppr_rpr):
        ppr_rpr.remove(child)
    if font_ascii:
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:ascii'), font_ascii)
        rf.set(qn('w:eastAsia'), 'Times New Roman')
        rf.set(qn('w:hAnsi'), font_ascii)
        rf.set(qn('w:cs'), font_ascii)
        ppr_rpr.append(rf)
    if color:
        c = OxmlElement('w:color'); c.set(qn('w:val'), color); ppr_rpr.append(c)
    if size:
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(size)); ppr_rpr.append(sz)
        szcs = OxmlElement('w:szCs'); szcs.set(qn('w:val'), str(size)); ppr_rpr.append(szcs)

    run = _make_run(text, bold=bold, size=size,
                    font_ascii=font_ascii, color=color)
    para._p.append(run)


def _set_alloc_amount_cell(cell, text):
    """
    Set allocation-table amount cell: font 9pt (sz=18), right-aligned.
    Matches correct doc: sz=114300 EMU = 9pt.
    """
    para = cell.paragraphs[0]
    _clear_para_runs(para)
    # Set right alignment
    ppr = para._p.find(qn('w:pPr'))
    if ppr is None:
        ppr = OxmlElement('w:pPr')
        para._p.insert(0, ppr)
    jc = ppr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        ppr.append(jc)
    jc.set(qn('w:val'), 'right')
    # Plain run with sz=18 (9pt)
    run = _make_run(text, size=18, font_ascii=None)
    para._p.append(run)


def _right_align_cell(cell):
    """Add right-alignment to every paragraph of a cell."""
    for para in cell.paragraphs:
        ppr = para._p.find(qn('w:pPr'))
        if ppr is None:
            ppr = OxmlElement('w:pPr')
            para._p.insert(0, ppr)
        jc = ppr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            ppr.append(jc)
        jc.set(qn('w:val'), 'right')

