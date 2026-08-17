"""
Labour Annexure (wage/attendance) document generator for THAYAMMA and
DEVAMMA. Leave days are randomly assigned with no overlap between workers.
"""
import docx
import random

from core.dates import get_days_in_month
from core.num_to_words import amount_in_words_title
from docx_engine.xml_helpers import _make_run, _clear_para_runs


def generate_labour_annexure(template_path, month, year,
                             tha_days, dev_days, wage):
    doc = docx.Document(template_path)
    days, m_idx = get_days_in_month(month, year)
    s_date = f"01/{m_idx:02d}/{year}"
    e_date = f"{days}/{m_idx:02d}/{year}"

    tha_total = tha_days * wage
    dev_total = dev_days * wage

    # Leave logic: random spread, guaranteed no overlap between workers
    tha_leaves     = days - tha_days
    dev_leaves     = days - dev_days
    all_month_days = list(range(1, days + 1))
    tha_leave_days = set(random.sample(all_month_days, tha_leaves)) if tha_leaves > 0 else set()
    remaining      = [d for d in all_month_days if d not in tha_leave_days]
    dev_leave_days = set(random.sample(remaining, min(dev_leaves, len(remaining)))) if dev_leaves > 0 else set()

    def _labour_run(para, text, bold=True, size=28):
        para._p.append(_make_run(text, bold=bold, size=size, font_ascii=None))

    # ── 1. PARAGRAPH REPLACEMENTS ────────────────────────────
    worker_order     = []
    total_paid_paras = []
    and_paid_paras   = []
    date_paras       = []

    # Scan ALL paragraphs: top-level AND inside table cells
    # (doc.paragraphs misses table cell paragraphs — that's why Total amount paid was empty)
    def _iter_all_paragraphs(doc):
        """Yield every paragraph in the document, including those inside table cells."""
        for p in doc.paragraphs:
            yield p
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    for p in _iter_all_paragraphs(doc):
        txt = p.text

        if "LABOUR NAME:" in txt:
            if "THAYAMMA" in txt.upper():
                worker_order.append("THAYAMMA")
                p.text = (f"         LABOUR NAME: Smt. THAYAMMA"
                          f"                                                              PERIOD: FROM {s_date} TO {e_date}")
            elif "DEVAMMA" in txt.upper():
                worker_order.append("DEVAMMA")
                p.text = (f"         LABOUR NAME: Smt. DEVAMMA"
                          f"                                                              PERIOD: FROM {s_date} TO {e_date}")

        elif "Total amount paid" in txt:
            total_paid_paras.append(p)

        elif "And paid Rs" in txt or "and paid Rs" in txt:
            and_paid_paras.append(p)

        elif "DATE:" in txt and (
                "Signature" in txt or "labourer" in txt.lower()
                or txt.strip().startswith("DATE:")):
            date_paras.append(p)

    # Fallback worker order
    if not worker_order:
        worker_order = ["THAYAMMA", "DEVAMMA"]
    elif len(worker_order) == 1:
        worker_order.append("DEVAMMA" if worker_order[0] == "THAYAMMA" else "THAYAMMA")

    # Fill "And paid Rs" paragraphs
    for idx, p in enumerate(and_paid_paras):
        cw = worker_order[idx] if idx < len(worker_order) else worker_order[-1]
        w_total = tha_total if cw == "THAYAMMA" else dev_total
        p.text = f"\t    And paid Rs {w_total}/- in presence of the following witness."

    # Fix DATE paragraphs
    for p in date_paras:
        txt = p.text
        colon_pos = txt.find("DATE:")
        if colon_pos >= 0:
            prefix = txt[:colon_pos + 5]   # up to and including "DATE:"
            suffix = ""
            if "Signature" in txt:
                sig_pos = txt.find("Signature")
                suffix = "\t\t\t\t\t" + txt[sig_pos:]
            p.text = f"{prefix} {e_date}{suffix}"

    # Fill "Total amount paid" paragraphs — 10pt font (size=20)
    for idx, p in enumerate(total_paid_paras):
        cw      = worker_order[idx] if idx < len(worker_order) else worker_order[-1]
        w_total = tha_total if cw == "THAYAMMA" else dev_total
        words   = amount_in_words_title(w_total)
        _clear_para_runs(p)
        _labour_run(p, "Total amount paid (in figures and words) = ", bold=True, size=20)
        _labour_run(p, f"{w_total}/- ",                               bold=True, size=20)
        _labour_run(p, words,                                          bold=True, size=20)

    # ── 2. FILL CALENDAR TABLES ──────────────────────────────
    # RULES:
    # • cols[2..11] = CLEANING AREAS — DO NOT TOUCH, keep template √/X as-is
    # • col[0] = date for that worked day only
    # • col[12] = wage amount for worked days; empty for unused template rows
    # • Leave days are SKIPPED — their template row keeps no date / no amount
    # • No-overlap: tha gets last N days off, dev gets the block just before that

    all_days        = list(range(1, days + 1))
    tha_worked_days = sorted(d for d in all_days if d not in tha_leave_days)
    dev_worked_days = sorted(d for d in all_days if d not in dev_leave_days)

    table_calendar_idx = 0
    for table in doc.tables:
        if len(table.rows) > 15:
            current_worker = worker_order[table_calendar_idx % 2]
            w_total      = tha_total if current_worker == "THAYAMMA" else dev_total
            worked_days  = tha_worked_days if current_worker == "THAYAMMA" else dev_worked_days

            work_ptr = 0
            for row in table.rows:
                cells = row.cells
                if len(cells) > 12 and "08" in cells[1].text:
                    if work_ptr < len(worked_days):
                        # Worked day: write date and wage; CLEANING AREAS cols untouched
                        cells[0].text  = f"{worked_days[work_ptr]:02d}/{m_idx:02d}/{year}"
                        cells[12].text = str(wage)
                        work_ptr += 1
                    else:
                        # Extra template row beyond worked days: clear date and amount only
                        cells[0].text  = ""
                        cells[12].text = ""
                        # cols[2..11] stay as template

                elif "TOTAL" in cells[0].text or (
                        len(cells) > 2 and "TOTAL" in cells[2].text):
                    cells[12].text = str(w_total)

            table_calendar_idx += 1

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


