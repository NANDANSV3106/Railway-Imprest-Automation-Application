"""
Annexure I (Master Bill) document generator.
Fills the fixed station imprest recoupment template via targeted
paragraph/table text replacement using the OOXML helpers in xml_helpers.py.
"""
import docx

from core.dates import get_days_in_month
from core.num_to_words import amount_in_words
from docx_engine.xml_helpers import (
    _make_run, _clear_para_runs, _set_para_spacing,
    _set_cell_run, _set_alloc_amount_cell, _right_align_cell,
)


def generate_annexure_1(template_path, month, year,
                        op_bal, atm_recoup, cash_total,
                        dev_total, tha_total, tot_exp, cl_bal,
                        sanc_amt, v1_amt, v2_amt, v1_date, v2_date,
                        v1_items=0, v2_items=0, num_vouchers=3, has_stationery=False):

    doc = docx.Document(template_path)
    days, m_idx = get_days_in_month(month, year)
    s_date   = f"01/{m_idx:02d}/{year}"
    e_date   = f"{days}/{m_idx:02d}/{year}"
    next_m   = 1 if m_idx == 12 else m_idx + 1
    next_y   = int(year) + 1 if m_idx == 12 else year
    next_date = f"01/{next_m:02d}/{next_y}"
    l_total  = dev_total + tha_total

    # ── 1. PARAGRAPH REPLACEMENTS ────────────────────────────
    for p in doc.paragraphs:
        txt = p.text

        if "No. Station Imprest" in txt:
            p.text = (f" No. Station Imprest /{month} {year}"
                      "                                                                                                        Station: SETTIHALLI")

        elif "Date:" in txt and "THROUGH" not in txt and "Dated:" not in txt:
            p.text = ("\t\t\t\t\t\t\t\t\t\t\t  Date:" + e_date)

        elif "Sub : Recoupment" in txt or ("Sub:" in txt and "Recoupment" in txt):
            p.text = f"\t\tSub: Recoupment of imprest bill for the period {s_date} to {e_date}"

        # ── Item 14: Vr. No. line ──
        elif ("14. Vr. No." in txt) or ("Vr. No.01" in txt) or \
             ("14." in txt and "Vr. No." in txt and "Date" in txt):
            _clear_para_runs(p)

            def ar(text, bold=False, size=20, underline=False):
                p._p.append(_make_run(text, bold=bold, size=size,
                                      underline=underline, font_ascii=None))

            ar("14. ",    bold=False, size=16)
            ar("Vr. No.", bold=False, size=20)
            ar("01",      bold=True,  size=20, underline=True)
            ar(" ",       bold=True,  size=20)
            ar("Date ",   bold=False, size=20)
            ar(v1_date,   bold=True,  size=20)
            ar(",",       bold=True,  size=20, underline=True)
            ar("Vr. No.", bold=False, size=20)
            ar("02",      bold=True,  size=20, underline=True)
            ar(" ",       bold=True,  size=20)
            ar("Date",    bold=False, size=20)
            ar(f" {v2_date if (v2_amt and v2_date) else e_date}", bold=True, size=20)
            ar(" ,Vr.No ", bold=False, size=20)
            ar("03",      bold=True,  size=20, underline=True)
            ar(" Date ",  bold=False, size=20)
            ar(e_date,    bold=True,  size=20)
            if num_vouchers == 4:
                ar(" ,Vr.No ", bold=False, size=20)
                ar("04",       bold=True,  size=20, underline=True)
                ar(" Date ",   bold=False, size=20)
                ar(e_date,     bold=True,  size=20)

        # ── "Accepted for Rs. /-(...)" paragraph ──
        elif "Accepted for Rs." in txt and "14." not in txt:
            _clear_para_runs(p)

            def ar2(text, bold=False, size=20):
                p._p.append(_make_run(text, bold=bold, size=size, font_ascii=None))

            words = amount_in_words(tot_exp)
            ar2("Accepted for Rs. ", bold=False, size=20)
            ar2(f"{tot_exp}/-",      bold=True,  size=20)
            ar2(words,               bold=True,  size=20)

        elif "For the month of" in txt:
            p.text = f"For the month of  {month}  {year}"

        elif "Imprest Account Of" in txt:
            p.text = (f"Imprest Account Of    SETTIHALLI    "
                      f"का   अग्रदाय   लेखा   महिन   {month}  {year}    के   लिए")

    # ── 2. TABLE ENGINE ──────────────────────────────────────
    def _set_t3_amount(cell, val):
        """Write a TABLE 3 amount cell: 12pt bold, right-aligned."""
        cell.text = val
        for _p in cell.paragraphs:
            for _r in _p.runs:
                _r.bold = True
                _r.font.size = docx.shared.Pt(12)
            _ppr = _p._p.find(qn('w:pPr'))
            if _ppr is None:
                _ppr = OxmlElement('w:pPr'); _p._p.insert(0, _ppr)
            _jc = _ppr.find(qn('w:jc'))
            if _jc is None:
                _jc = OxmlElement('w:jc'); _ppr.append(_jc)
            _jc.set(qn('w:val'), 'right')
            _rpr = _ppr.find(qn('w:rPr'))
            if _rpr is None:
                _rpr = OxmlElement('w:rPr'); _ppr.append(_rpr)
            if _rpr.find(qn('w:b')) is None:
                _rpr.append(OxmlElement('w:b'))
            _sz = _rpr.find(qn('w:sz'))
            if _sz is None:
                _sz = OxmlElement('w:sz'); _rpr.append(_sz)
            _sz.set(qn('w:val'), '24')   # 12pt = sz 24


    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            texts    = [c.text.strip() for c in row.cells]
            full_txt = " ".join(texts)

            # TABLE 0: Balance matrix
            if "Opening Balance" in full_txt and "Now recoupment" in full_txt \
                    and len(row.cells) >= 4:
                row.cells[0].text = f"Opening Balance as on {s_date}"
                row.cells[1].text = str(op_bal)
                # Preserve bold on Rs. column
                for run in row.cells[1].paragraphs[0].runs:
                    run.bold = True
                _right_align_cell(row.cells[1])
                row.cells[2].text = f"Now recoupment Bill No. NIL,\nDated: {e_date}"
                row.cells[3].text = str(tot_exp)
                for run in row.cells[3].paragraphs[0].runs:
                    run.bold = True
                _right_align_cell(row.cells[3])

            elif "Last recoupment" in full_txt and "Closing" not in full_txt \
                    and len(row.cells) >= 2:
                # FIX: no comma
                row.cells[0].text = f"Last recoupment Bill No.NIL Dated {month} {year}"
                row.cells[1].text = str(atm_recoup)
                for run in row.cells[1].paragraphs[0].runs:
                    run.bold = True
                _right_align_cell(row.cells[1])

            elif "Closing Balance" in full_txt and len(row.cells) >= 4:
                row.cells[2].text = f"Closing Balance as on {e_date}"
                row.cells[3].text = str(cl_bal)
                for run in row.cells[3].paragraphs[0].runs:
                    run.bold = True
                _right_align_cell(row.cells[3])

            elif "Sanctioned Amount" in full_txt and len(row.cells) >= 4:
                row.cells[1].text = str(sanc_amt)
                row.cells[3].text = str(sanc_amt)
                for c in [row.cells[1], row.cells[3]]:
                    for run in c.paragraphs[0].runs:
                        run.bold = True
                    _right_align_cell(c)

            # TABLE 2: 7-column summary — match by Sl.No (works for 3V & 4V)
            elif "09.295" in full_txt and len(row.cells) == 7:
                sl = texts[0]

                def _set_t2_cell(cell, val, bold=True, size=20):
                    """Write text into a Table-2 cell at 10pt (size=20), optional bold, right-aligned."""
                    cell.text = val
                    for _p in cell.paragraphs:
                        for _r in _p.runs:
                            _r.bold = bold
                            _r.font.size = docx.shared.Pt(10)
                        # inject bold into paragraph mark so empty cells also bold
                        _ppr = _p._p.find(qn('w:pPr'))
                        if _ppr is None:
                            _ppr = OxmlElement('w:pPr'); _p._p.insert(0, _ppr)
                        _jc = _ppr.find(qn('w:jc'))
                        if _jc is None:
                            _jc = OxmlElement('w:jc'); _ppr.append(_jc)
                        _jc.set(qn('w:val'), 'right')
                        _rpr = _ppr.find(qn('w:rPr'))
                        if _rpr is None:
                            _rpr = OxmlElement('w:rPr'); _ppr.append(_rpr)
                        _b = _rpr.find(qn('w:b'))
                        if bold and _b is None:
                            _rpr.append(OxmlElement('w:b'))
                        elif not bold and _b is not None:
                            _rpr.remove(_b)
                        _sz = _rpr.find(qn('w:sz'))
                        if _sz is None:
                            _sz = OxmlElement('w:sz'); _rpr.append(_sz)
                        _sz.set(qn('w:val'), '20')   # 10pt = sz 20

                def _fmt_items(n):
                    """Format item count: single digit gets zero-padded to 2 chars."""
                    return f"{int(n):02d}" if n else ""

                if sl in ("01", "1"):
                    row.cells[1].text = v1_date
                    _set_t2_cell(row.cells[5], _fmt_items(v1_items), bold=True)
                    _set_t2_cell(row.cells[6], str(v1_amt), bold=True)

                elif sl in ("02", "2"):
                    if num_vouchers == 3:
                        row.cells[1].text = e_date
                        _set_t2_cell(row.cells[6], str(dev_total), bold=True)
                    else:
                        row.cells[1].text = v2_date
                        _set_t2_cell(row.cells[5], _fmt_items(v2_items), bold=True)
                        _set_t2_cell(row.cells[6], str(v2_amt), bold=True)

                elif sl in ("03", "3"):
                    if num_vouchers == 3:
                        row.cells[1].text = e_date
                        _set_t2_cell(row.cells[6], str(tha_total), bold=True)
                    else:
                        row.cells[1].text = e_date
                        _set_t2_cell(row.cells[6], str(dev_total), bold=True)

                elif sl in ("04", "4"):
                    row.cells[1].text = e_date
                    _set_t2_cell(row.cells[6], str(tha_total), bold=True)

                # Set 10pt on all non-amount cells in this data row
                for _ci, _cell in enumerate(row.cells):
                    if _ci not in (5, 6):
                        for _p in _cell.paragraphs:
                            for _r in _p.runs:
                                _r.font.size = docx.shared.Pt(10)

            elif texts[0] == "TOTAL" and len(row.cells) == 7:
                row.cells[6].text = str(tot_exp)
                # TOTAL row: entire row bold, 14pt
                for _cell in row.cells:
                    for _p in _cell.paragraphs:
                        for _r in _p.runs:
                            _r.bold = True
                            _r.font.size = docx.shared.Pt(12)
                        _ppr = _p._p.find(qn('w:pPr'))
                        if _ppr is None:
                            _ppr = OxmlElement('w:pPr'); _p._p.insert(0, _ppr)
                        _rpr = _ppr.find(qn('w:rPr'))
                        if _rpr is None:
                            _rpr = OxmlElement('w:rPr'); _ppr.append(_rpr)
                        if _rpr.find(qn('w:b')) is None:
                            _rpr.append(OxmlElement('w:b'))
                        _sz = _rpr.find(qn('w:sz'))
                        if _sz is None:
                            _sz = OxmlElement('w:sz'); _rpr.append(_sz)
                        _sz.set(qn('w:val'), '20')   # 14pt = sz 20
                _right_align_cell(row.cells[6])

            # TABLE 3: G.I. Ledger
            elif "OPENING BALANCE" in full_txt and len(row.cells) >= 6:
                row.cells[0].text = s_date
                row.cells[3].text = str(op_bal)
                _right_align_cell(row.cells[3])

            elif "Amount received through ATM" in full_txt and len(row.cells) >= 4:
                row.cells[3].text = str(atm_recoup)
                _right_align_cell(row.cells[3])

            elif "Consumable Items" in full_txt and len(row.cells) >= 6:
                vno = texts[1]
                if vno in ("1", "01"):
                    row.cells[0].text = v1_date
                    _set_t3_amount(row.cells[5], str(v1_amt))
                elif vno in ("2", "02") and num_vouchers == 4:
                    row.cells[0].text = v2_date
                    _set_t3_amount(row.cells[5], str(v2_amt))

            # TABLE 3 Labour rows — 3V: vouchers 2&3; 4V: vouchers 3&4
            elif "Amount paid to cleaning labour" in full_txt and len(row.cells) >= 6:
                vno = texts[1]
                if num_vouchers == 3:
                    if vno in ("2", "02"):
                        row.cells[0].text = e_date
                        _set_t3_amount(row.cells[5], str(dev_total))
                    elif vno in ("3", "03"):
                        row.cells[0].text = e_date
                        _set_t3_amount(row.cells[5], str(tha_total))
                else:
                    if vno in ("3", "03"):
                        row.cells[0].text = e_date
                        _set_t3_amount(row.cells[5], str(dev_total))
                    elif vno in ("4", "04"):
                        row.cells[0].text = e_date
                        _set_t3_amount(row.cells[5], str(tha_total))

            elif "TOTAL" in full_txt and "OPENING" not in full_txt \
                    and len(row.cells) == 8:
                row.cells[3].text = str(sanc_amt)
                row.cells[5].text = str(tot_exp)
                # TOTAL row: bold + 14pt entire row + right-align amount cells
                for _cell in row.cells:
                    for _p in _cell.paragraphs:
                        for _r in _p.runs:
                            _r.bold = True
                            _r.font.size = docx.shared.Pt(14)
                        _ppr = _p._p.find(qn('w:pPr'))
                        if _ppr is None:
                            _ppr = OxmlElement('w:pPr'); _p._p.insert(0, _ppr)
                        _rpr = _ppr.find(qn('w:rPr'))
                        if _rpr is None:
                            _rpr = OxmlElement('w:rPr'); _ppr.append(_rpr)
                        if _rpr.find(qn('w:b')) is None:
                            _rpr.append(OxmlElement('w:b'))
                        _sz = _rpr.find(qn('w:sz'))
                        if _sz is None:
                            _sz = OxmlElement('w:sz'); _rpr.append(_sz)
                        _sz.set(qn('w:val'), '22')   # 14pt
                _right_align_cell(row.cells[3])
                _right_align_cell(row.cells[5])

            # TABLE 4: Certified section
            elif "BALANCE AMOUNT ON HAND" in full_txt:
                _set_cell_run(
                    row.cells[0],
                    f"BALANCE AMOUNT ON HAND AS ON {e_date} IS RS. {cl_bal}/-",
                    size=28, spacing_after="0", spacing_line="240"
                )

            elif "2. Certified that the amount paid" in full_txt:
                # Amount Rs.X/- must be bold; rest plain
                cell = row.cells[0]
                para = cell.paragraphs[0]
                _clear_para_runs(para)
                _set_para_spacing(para, after="0", line="240")
                def _a1_run(text, bold=False):
                    para._p.append(_make_run(text, bold=bold, size=28, font_ascii=None))
                _a1_run("2. Certified that the amount paid and voucher accepted for Rs. ")
                _a1_run(f"{tot_exp}/-", bold=True)

            elif "3. Certified that the amount on hand" in full_txt:
                # Amount Rs.X/- must be bold; rest plain
                cell = row.cells[0]
                para = cell.paragraphs[0]
                _clear_para_runs(para)
                _set_para_spacing(para, after="0", line="240")
                def _a1_run3(text, bold=False):
                    para._p.append(_make_run(text, bold=bold, size=28, font_ascii=None))
                _a1_run3(f"3. Certified that the amount on hand as on {next_date} is Rs. ")
                _a1_run3(f"{cl_bal}/-", bold=True)

    # ── TABLE 4 — words rows (rows 3 and 5) ─────────────────
    # FIX: use _set_cell_run so font=14pt, spacing tight — no huge gaps
    if len(doc.tables) >= 5:
        t4 = doc.tables[4]
        if len(t4.rows) > 5:
            _set_cell_run(
                t4.rows[3].cells[0],
                amount_in_words(tot_exp),
                size=28,bold=True, spacing_after="0", spacing_line="240"
            )
            _set_cell_run(
                t4.rows[5].cells[0],
                amount_in_words(cl_bal),
                size=28,bold=True,spacing_after="0", spacing_line="240"
            )

    # ── 3. NESTED ALLOCATION TABLE ───────────────────────────
    # FIX: allocation amounts get font size 9pt (sz=18), right-aligned
    if len(doc.tables) > 1:
        alloc_wrapper = doc.tables[1]
        if alloc_wrapper.rows and alloc_wrapper.rows[0].cells:
            cell = alloc_wrapper.rows[0].cells[0]
            if cell.tables:
                nested = cell.tables[0]
                for row in nested.rows:
                    texts    = [c.text.strip() for c in row.cells]
                    full_txt = " ".join(texts)
                    if texts[0] in ("01", "1") and "09.295" in full_txt:
                        _set_alloc_amount_cell(row.cells[2], str(cash_total))
                        if has_stationery:
                            row.cells[4].text = "Consumable items bills & Stationery item bills"
                        else:
                            row.cells[4].text = "Consumable items billss"
                    elif texts[0] in ("02", "2") and "09.295" in full_txt:
                        _set_alloc_amount_cell(row.cells[2], str(l_total))
                        row.cells[4].text = "Cleaning labour bills"
                       
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


